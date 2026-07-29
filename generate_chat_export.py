import json
import os
import sys
import subprocess

# Ensure python-docx and reportlab are installed
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "reportlab"])
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

# Register Korean font for PDF (Malgun Gothic)
font_path = r"C:\Windows\Fonts\malgun.ttf"
font_bold_path = r"C:\Windows\Fonts\malgunbd.ttf"

has_korean_font = False
if os.path.exists(font_path):
    pdfmetrics.registerFont(TTFont('MalgunGothic', font_path))
    if os.path.exists(font_bold_path):
        pdfmetrics.registerFont(TTFont('MalgunGothic-Bold', font_bold_path))
    else:
        pdfmetrics.registerFont(TTFont('MalgunGothic-Bold', font_path))
    has_korean_font = True

# Path to transcript
transcript_path = r"C:\Users\user\.gemini\antigravity-ide\brain\63e68240-ffee-4b92-8b3d-d62e69fbd22a\.system_generated\logs\transcript_full.jsonl"
if not os.path.exists(transcript_path):
    transcript_path = r"C:\Users\user\.gemini\antigravity-ide\brain\63e68240-ffee-4b92-8b3d-d62e69fbd22a\.system_generated\logs\transcript.jsonl"

print(f"Reading transcript from: {transcript_path}")

history = []

with open(transcript_path, "r", encoding="utf-8") as f:
    for line in f:
        line = line.strip()
        if not line:
            continue
        try:
            data = json.loads(line)
        except Exception:
            continue
        
        step_type = data.get("type")
        source = data.get("source")
        
        if step_type == "USER_INPUT" or source == "USER_EXPLICIT":
            content = data.get("content", "")
            if isinstance(content, dict):
                text = content.get("text", "")
            elif isinstance(content, list):
                text = " ".join([str(item) for item in content])
            else:
                text = str(content)
            
            # Clean XML-like tags if present
            if "<USER_REQUEST>" in text:
                start = text.find("<USER_REQUEST>") + len("<USER_REQUEST>")
                end = text.find("</USER_REQUEST>")
                if end != -1:
                    text = text[start:end].strip()
                else:
                    text = text[start:].strip()
            
            if text and not text.startswith("{{ CHECKPOINT"):
                history.append({"role": "사용자 (User)", "text": text})
                
        elif step_type == "PLANNER_RESPONSE" or source == "MODEL":
            content = data.get("content", "")
            text_parts = []
            if isinstance(content, dict):
                text_parts.append(content.get("text", ""))
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, dict) and item.get("type") == "text":
                        text_parts.append(item.get("text", ""))
                    elif isinstance(item, str):
                        text_parts.append(item)
            else:
                text_parts.append(str(content))
            
            full_text = "\n".join(text_parts).strip()
            if full_text:
                history.append({"role": "AI 도우미 (Antigravity)", "text": full_text})

print(f"Total extracted Q&A entries: {len(history)}")

# Output filenames
output_docx = r"c:\Users\user\Desktop\indivi\AI수익화\결제 수익화런칭\dream-teller\AI_Dream_Teller_전체_대화_록.docx"
output_pdf = r"c:\Users\user\Desktop\indivi\AI수익화\결제 수익화런칭\dream-teller\AI_Dream_Teller_전체_대화_록.pdf"

# 1. Create DOCX
doc = docx.Document()
doc.add_heading('AI Dream Teller 개발 프로젝트 전체 대화 기록', 0)

for idx, entry in enumerate(history):
    role = entry['role']
    text = entry['text']
    
    p_role = doc.add_paragraph()
    run_role = p_role.add_run(f"[{idx+1}] {role}")
    run_role.bold = True
    if "사용자" in role:
        run_role.font.color.rgb = docx.shared.RGBColor(139, 92, 246) # Purple
    else:
        run_role.font.color.rgb = docx.shared.RGBColor(16, 185, 129) # Emerald Green
        
    p_text = doc.add_paragraph(text)
    doc.add_paragraph("─" * 40)

doc.save(output_docx)
print(f"Successfully generated DOCX: {output_docx}")

# 2. Create PDF
styles = getSampleStyleSheet()

pdf_font = 'MalgunGothic' if has_korean_font else 'Helvetica'
pdf_bold_font = 'MalgunGothic-Bold' if has_korean_font else 'Helvetica-Bold'

style_title = ParagraphStyle(
    'TitleStyle',
    parent=styles['Heading1'],
    fontName=pdf_bold_font,
    fontSize=18,
    leading=22,
    textColor=colors.HexColor("#6B21A8"),
    spaceAfter=15
)

style_user_role = ParagraphStyle(
    'UserRole',
    parent=styles['Normal'],
    fontName=pdf_bold_font,
    fontSize=12,
    leading=16,
    textColor=colors.HexColor("#7C3AED"),
    spaceBefore=10,
    spaceAfter=5
)

style_ai_role = ParagraphStyle(
    'AIRole',
    parent=styles['Normal'],
    fontName=pdf_bold_font,
    fontSize=12,
    leading=16,
    textColor=colors.HexColor("#059669"),
    spaceBefore=10,
    spaceAfter=5
)

style_body = ParagraphStyle(
    'BodyText',
    parent=styles['Normal'],
    fontName=pdf_font,
    fontSize=10,
    leading=14,
    textColor=colors.HexColor("#1F2937"),
    spaceAfter=10
)

doc_pdf = SimpleDocTemplate(
    output_pdf,
    pagesize=letter,
    rightMargin=40,
    leftMargin=40,
    topMargin=40,
    bottomMargin=40
)

story = [
    Paragraph("AI Dream Teller 개발 프로젝트 전체 대화 기록", style_title),
    HRFlowable(width="100%", thickness=1, color=colors.HexColor("#DDD6FE"), spaceAfter=15)
]

for idx, entry in enumerate(history):
    role = entry['role']
    text = entry['text']
    
    role_style = style_user_role if "사용자" in role else style_ai_role
    story.append(Paragraph(f"[{idx+1}] {role}", role_style))
    
    # Format text for PDF HTML paragraph
    formatted_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
    story.append(Paragraph(formatted_text, style_body))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#E5E7EB"), spaceAfter=10))

doc_pdf.build(story)
print(f"Successfully generated PDF: {output_pdf}")
