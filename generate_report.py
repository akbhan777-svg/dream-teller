import os
import json
import re
import shutil
import datetime

# 1. 한국어 맑은 고딕 폰트 경로 (Windows 표준)
FONT_PATH = "C:\\Windows\\Fonts\\malgun.ttf"

CONVERSATION_IDS = [
    "4315ab1d-5e8d-40bc-a4da-502b2128c005", # 이전 대화 세션
    "bb6a8237-f48e-4f04-bcc5-766550877bba"  # 현재 대화 세션
]

BASE_BRAIN_PATH = r"C:\Users\user\.gemini\antigravity-ide\brain"

def parse_all_transcripts():
    dialogues = []
    
    for conv_id in CONVERSATION_IDS:
        conv_dir = os.path.join(BASE_BRAIN_PATH, conv_id, ".system_generated", "logs")
        transcript_path = os.path.join(conv_dir, "transcript_full.jsonl")
        if not os.path.exists(transcript_path):
            transcript_path = os.path.join(conv_dir, "transcript.jsonl")
        
        if not os.path.exists(transcript_path):
            print(f"경로를 찾을 수 없음: {transcript_path}")
            continue

        current_user_request = None

        with open(transcript_path, 'r', encoding='utf-8') as f:
            for line in f:
                try:
                    data = json.loads(line)
                    # 유저 입력 추출
                    if data.get("type") == "USER_INPUT":
                        content = data.get("content", "")
                        # XML 태그 제거 (<USER_REQUEST> 등)
                        clean_content = re.sub(r'<[^>]+>', '', content).strip()
                        # 추가 메타데이터 제거
                        clean_content = clean_content.split("<ADDITIONAL_METADATA>")[0].strip()
                        clean_content = clean_content.split("<USER_SETTINGS_CHANGE>")[0].strip()
                        
                        if clean_content and "CHECKPOINT" not in clean_content and "Conversation History" not in clean_content:
                            current_user_request = clean_content
                    
                    # 에이전트 답변 추출
                    elif data.get("type") == "PLANNER_RESPONSE" and current_user_request:
                        response = data.get("content", "")
                        if response and response.strip():
                            # 마크다운 파일 링크 가독성 개선
                            clean_response = re.sub(r'\[([^\]]+)\]\(file:///[^\)]+\)', r'\1', response)
                            dialogues.append({
                                "session": conv_id[:8],
                                "question": current_user_request,
                                "answer": clean_response.strip()
                            })
                            current_user_request = None
                except Exception:
                    continue
                    
    # 중복 제거 및 정제
    unique_dialogues = []
    seen_questions = set()
    for item in dialogues:
        q_key = item["question"][:50]
        if q_key not in seen_questions:
            seen_questions.add(q_key)
            unique_dialogues.append(item)
            
    return unique_dialogues

def generate_docx(dialogues, output_path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
    except ImportError:
        print("python-docx 라이브러리가 필요합니다.")
        return False

    doc = Document()
    
    # 기본 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10)
    
    # 여백 설정
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # 제목 추가
    title = doc.add_paragraph()
    title_run = title.add_run("Dream Teller 전체 개발 대화 및 질의응답 리포트 (전체)")
    title_run.font.name = 'Malgun Gothic'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(79, 70, 229) # Purple
    
    today_str = datetime.datetime.now().strftime("%Y년 %m월 %d일 %H:%M")
    p_meta = doc.add_paragraph(f"작성일시: {today_str} | 총 질의응답 건수: {len(dialogues)}건\n")
    p_meta.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    for idx, item in enumerate(dialogues, 1):
        # 질문 섹션
        q_p = doc.add_paragraph()
        q_run = q_p.add_run(f"Q{idx}. [질문] {item['question']}")
        q_run.font.bold = True
        q_run.font.size = Pt(11.5)
        q_run.font.color.rgb = RGBColor(219, 39, 119) # Pink
        
        # 답변 섹션
        a_p = doc.add_paragraph()
        a_run = a_p.add_run(f"[답변]\n{item['answer']}\n")
        a_run.font.size = Pt(10)
        a_run.font.color.rgb = RGBColor(31, 41, 55)
        
        # 구분선 추가
        div_p = doc.add_paragraph("-" * 85)
        div_p.runs[0].font.color.rgb = RGBColor(209, 213, 219)
        
    doc.save(output_path)
    print(f"Word 파일 생성 성공: {output_path}")
    return True

def generate_pdf(dialogues, output_path):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttcharts import TTFont
    except ImportError as e:
        print(f"reportlab 불러오기 실패: {e}")
        return False

    if not os.path.exists(FONT_PATH):
        print(f"폰트 없음: {FONT_PATH}")
        return False

    pdfmetrics.registerFont(TTFont('Malgun', FONT_PATH))

    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36
    )

    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'PDFTitle',
        parent=styles['Heading1'],
        fontName='Malgun',
        fontSize=16,
        textColor=colors.HexColor('#4f46e5'),
        spaceAfter=10
    )
    
    q_style = ParagraphStyle(
        'PDFQuestion',
        parent=styles['Heading3'],
        fontName='Malgun',
        fontSize=10.5,
        textColor=colors.HexColor('#db2777'),
        spaceBefore=10,
        spaceAfter=6,
        leading=14
    )

    a_style = ParagraphStyle(
        'PDFAnswer',
        parent=styles['Normal'],
        fontName='Malgun',
        fontSize=9,
        textColor=colors.HexColor('#1f2937'),
        spaceAfter=10,
        leading=13
    )

    meta_style = ParagraphStyle(
        'PDFMeta',
        parent=styles['Normal'],
        fontName='Malgun',
        fontSize=8.5,
        textColor=colors.HexColor('#6b7280'),
        spaceAfter=10
    )

    elements = []
    elements.append(Paragraph("Dream Teller 전체 개발 대화 및 질의응답 리포트", title_style))
    today_str = datetime.datetime.now().strftime("%Y년 %m월 %d일 %H:%M")
    elements.append(Paragraph(f"작성일시: {today_str} | 총 질의응답: {len(dialogues)}건", meta_style))
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e5e7eb'), spaceAfter=10))

    for idx, item in enumerate(dialogues, 1):
        q_text = f"Q{idx}. [질문] {item['question']}".replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
        elements.append(Paragraph(q_text, q_style))
        
        a_text = f"[답변]<br/>{item['answer']}".replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
        elements.append(Paragraph(a_text, a_style))
        
        elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#f3f4f6'), spaceBefore=8, spaceAfter=8))

    doc.build(elements)
    print(f"PDF 파일 생성 성공: {output_path}")
    return True

def generate_html(dialogues, output_path):
    today_str = datetime.datetime.now().strftime("%Y년 %m월 %d일 %H:%M")
    
    html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Dream Teller 전체 개발 대화 및 질의응답 리포트</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;700;900&display=swap');
        body {{
            font-family: 'Noto Sans KR', 'Malgun Gothic', sans-serif;
            background-color: #f9fafb;
            color: #1f2937;
            line-height: 1.7;
            padding: 40px 20px;
            max-width: 900px;
            margin: 0 auto;
        }}
        .header {{
            border-bottom: 2px solid #e5e7eb;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        h1 {{
            color: #4f46e5;
            font-size: 26px;
            font-weight: 900;
            margin: 0 0 10px 0;
        }}
        .meta {{
            font-size: 14px;
            color: #6b7280;
        }}
        .item {{
            margin-bottom: 35px;
            padding-bottom: 20px;
            border-bottom: 1px solid #f3f4f6;
        }}
        .question {{
            color: #db2777;
            font-size: 15px;
            font-weight: 700;
            margin-bottom: 12px;
            background-color: #fdf2f8;
            padding: 12px 18px;
            border-left: 4px solid #db2777;
            border-radius: 4px;
        }}
        .answer {{
            font-size: 14px;
            white-space: pre-wrap;
            padding: 5px 10px;
            color: #374151;
        }}
        @media print {{
            body {{
                background-color: #ffffff;
                padding: 0;
            }}
            .question {{
                background-color: #fdf2f8 !important;
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }}
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Dream Teller 전체 개발 대화 및 질의응답 리포트</h1>
        <div class="meta">작성일시: {today_str} | 총 질의응답: {len(dialogues)}건</div>
    </div>
"""
    for idx, item in enumerate(dialogues, 1):
        q_text = item['question'].replace('<', '&lt;').replace('>', '&gt;')
        a_text = item['answer'].replace('<', '&lt;').replace('>', '&gt;')
        html_content += f"""
    <div class="item">
        <div class="question">Q{idx}. [질문] {q_text}</div>
        <div class="answer">[답변]<br/>{a_text}</div>
    </div>
"""
    html_content += """
</body>
</html>
"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    print(f"HTML 파일 생성 성공: {output_path}")
    return True

def main():
    print("모든 트랜스크립트 파싱 시작...")
    dialogues = parse_all_transcripts()
    
    if not dialogues:
        print("대화 내역을 가져오지 못했습니다.")
        return
        
    print(f"총 {len(dialogues)}개의 질의응답 추출 완료.")
    
    # 1. 루트 경로 파일들
    root_docx = "dream_teller_chat_report.docx"
    root_pdf = "dream_teller_chat_report.pdf"
    root_html = "dream_teller_chat_report.html"
    
    # 2. public 경로 (웹 브라우저 다운로드용)
    public_dir = os.path.join(os.getcwd(), "public")
    os.makedirs(public_dir, exist_ok=True)
    
    pub_docx = os.path.join(public_dir, "dream_teller_chat_report.docx")
    pub_pdf = os.path.join(public_dir, "dream_teller_chat_report.pdf")
    pub_html = os.path.join(public_dir, "dream_teller_chat_report.html")

    # 생성
    generate_docx(dialogues, root_docx)
    generate_pdf(dialogues, root_pdf)
    generate_html(dialogues, root_html)
    
    # Public 폴더에 복사
    shutil.copy(root_docx, pub_docx)
    if os.path.exists(root_pdf):
        shutil.copy(root_pdf, pub_pdf)
    shutil.copy(root_html, pub_html)

    print("복사 완료: public 디렉토리")

if __name__ == "__main__":
    main()
